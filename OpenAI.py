
import openai
import speech_recognition as sr
from gtts import gTTS
import pygame
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Set up OpenAI API credentials
openai.api_key = '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'

class Assistant:
    def __init__(self):
        self.document = Document()

    def speech_to_text(self):
        r = sr.Recognizer()
        with sr.Microphone() as source:
            print("Listening...")
            audio = r.listen(source)

        try:
            text = r.recognize_google(audio)
            print("User:", text)
            return text
        except sr.UnknownValueError:
            print("Sorry, I couldn't understand that.")
        except sr.RequestError as e:
            print("Could not request results from the speech recognition service; {0}".format(e))

    def text_to_speech(self, text):
        tts = gTTS(text=text, lang='en')
        tts.save('assistant_response.mp3')
        pygame.mixer.init()
        pygame.mixer.music.load('assistant_response.mp3')
        pygame.mixer.music.play()

    def chat_with_assistant(self, prompt):
        try:
            response = openai.Completion.create(
                engine='text-davinci-003',
                prompt=prompt,
                max_tokens=200
            )
            assistant_response = response.choices[0].text.strip()
            return assistant_response
        except Exception as e:
            print("An error occurred:", str(e))
            return None

    def handle_command(self, command):
        if command.lower() == "help":
            return "Sure, how can I assist you?"
        elif command.lower() == "time":
            import datetime
            current_time = datetime.datetime.now().strftime("%H:%M")
            return f"The current time is {current_time}."
        else:
            return None

    def run_conversation(self):
        print("AI Desktop Assistant:")
        print("You can start chatting. Speak or type 'exit' to quit.")

        while True:
            user_input = self.speech_to_text()

            if user_input is None:
                continue

            if user_input.lower() == 'exit':
                break

            # Check for specific commands
            command_response = self.handle_command(user_input)
            if command_response:
                print("Assistant:", command_response)
                self.text_to_speech(command_response)

                # Add the question and answer to the Word document
                self.document.add_heading('Question:', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                self.document.add_paragraph(user_input)
                self.document.add_heading('Answer:', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                self.document.add_paragraph(command_response)

                continue

            # Generate response from the AI assistant
            prompt = "User: " + user_input + "\nAssistant:"
            assistant_response = self.chat_with_assistant(prompt)

            if assistant_response:
                print("Assistant:", assistant_response)
                self.text_to_speech(assistant_response)

                # Add the question and answer to the Word document
                self.document.add_heading('Question:', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                self.document.add_paragraph(user_input)
                self.document.add_heading('Answer:', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                self.document.add_paragraph(assistant_response)

        # Save the Word document
        self.document.save('Conservation.docx')

        print("AI Desktop Assistant: Conversation ended.")




